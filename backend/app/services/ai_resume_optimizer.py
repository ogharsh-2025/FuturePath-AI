import re
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class AIResumeOptimizer:
    COMMON_TECH_KEYWORDS = [
        "python", "javascript", "typescript", "java", "c++", "go", "rust", "ruby", "php", "sql", "nosql",
        "react", "angular", "vue", "next.js", "node.js", "fastapi", "django", "flask", "spring", "express",
        "docker", "kubernetes", "aws", "gcp", "azure", "ci/cd", "git", "github", "terraform", "ansible",
        "machine learning", "deep learning", "nlp", "computer vision", "tensorflow", "pytorch", "scikit-learn",
        "postgresql", "mongodb", "redis", "mysql", "elasticsearch", "graphql", "rest api", "microservices",
        "html", "css", "tailwind", "bootstrap", "linux", "bash", "agile", "scrum", "jira"
    ]

    WEAK_VERBS_MAP = {
        "worked on": "engineered",
        "helped": "facilitated",
        "assisted": "collaborated on",
        "made": "architected",
        "did": "executed",
        "wrote": "implemented",
        "managed": "orchestrated",
        "led": "spearheaded",
        "handled": "managed",
        "responsible for": "drove",
        "created": "pioneered",
        "changed": "refactored",
        "looked at": "analyzed",
        "fixed": "resolved"
    }

    ACTION_VERBS = [
        "engineered", "architected", "spearheaded", "orchestrated", "implemented", "optimized",
        "formulated", "designed", "deployed", "streamlined", "facilitated", "drove", "pioneered",
        "collaborated", "refactored", "analyzed", "resolved", "developed", "built", "managed"
    ]

    @classmethod
    def analyze_resume_text(cls, text: str) -> dict:
        text_lower = text.lower()
        
        # 1. ATS Score calculation (out of 100)
        # Check standard sections
        sections = {
            "experience": ["experience", "work history", "employment", "professional background"],
            "education": ["education", "academic", "university", "college"],
            "skills": ["skills", "technologies", "technical expertise", "core competencies"],
            "projects": ["projects", "personal projects", "open source"]
        }
        
        ats_score = 50
        found_sections = []
        for sec_name, keywords in sections.items():
            found = False
            for kw in keywords:
                if kw in text_lower:
                    found = True
                    break
            if found:
                ats_score += 10
                found_sections.append(sec_name)
                
        # Check for contact info
        has_email = "@" in text_lower
        has_phone = bool(re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b|\b\+?\d{1,3}[-.\s]?\d{9,11}\b', text))
        has_github = "github.com" in text_lower
        has_linkedin = "linkedin.com" in text_lower
        
        if has_email: ats_score += 3
        if has_phone: ats_score += 3
        if has_github: ats_score += 2
        if has_linkedin: ats_score += 2
        
        ats_score = min(100, ats_score)

        # 2. Readability Score (heuristic based on average sentence and word length)
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        avg_sentence_len = len(words) / max(1, len(sentences))
        avg_word_len = sum(len(w) for w in words) / max(1, len(words))
        
        # Flesch-like heuristic: shorter sentences & shorter words -> higher readability
        readability_score = int(100 - (avg_sentence_len * 1.5) - (avg_word_len * 10))
        readability_score = max(30, min(100, readability_score))

        # 3. Formatting Score
        # Check for bullet points and spacing
        bullet_points = text.count("•") + text.count("- ") + text.count("* ")
        formatting_score = 60
        if bullet_points > 8:
            formatting_score += 20
        if len(text) > 500 and len(text) < 4000:
            formatting_score += 20
        formatting_score = min(100, formatting_score)

        # 4. Keyword Score
        matched_keywords = []
        missing_keywords = []
        for kw in cls.COMMON_TECH_KEYWORDS:
            if re.search(r'\b' + re.escape(kw) + r'\b', text_lower):
                matched_keywords.append(kw)
            else:
                missing_keywords.append(kw)
                
        keyword_score = int((len(matched_keywords) / max(1, len(matched_keywords) + min(15, len(missing_keywords)))) * 100)
        keyword_score = max(40, min(100, keyword_score))

        # 5. Project Quality Score & suggestions
        # Check if project descriptions use numbers, metrics, or strong action verbs
        has_metrics = bool(re.search(r'\b\d+%\b|\b\$\d+|\b\d+\s+users\b|\b\d+x\b|\breduced\b|\bincreased\b', text_lower))
        project_score = 65
        if has_metrics:
            project_score += 20
        
        action_verb_count = 0
        for verb in cls.ACTION_VERBS:
            if verb in text_lower:
                action_verb_count += 1
        if action_verb_count >= 5:
            project_score += 15
        project_score = min(100, project_score)

        # Suggestions lists
        weak_bullets = []
        bullet_matches = re.findall(r'(?:[•\-*]|\b\d+\.)\s*(.*?)(?=[•\-*]|\b\d+\.|\n\n|$)', text)
        for bullet in bullet_matches:
            bullet = bullet.strip()
            if not bullet:
                continue
            # Weak if short or contains weak verbs or lacks metrics
            has_strong_verb = any(verb in bullet.lower() for verb in cls.ACTION_VERBS)
            has_metric = bool(re.search(r'\d+%\s*|\$\s*\d+|\b\d+\s+percent\b|\b\d+\s+(?:hours|days|weeks|months|years|users|servers|records)\b', bullet.lower()))
            if len(bullet) < 150 and (not has_strong_verb or not has_metric):
                weak_bullets.append(bullet)

        # Better action verbs suggestions
        suggested_verb_replacements = {}
        for weak_verb, strong_verb in cls.WEAK_VERBS_MAP.items():
            if re.search(r'\b' + re.escape(weak_verb) + r'\b', text_lower):
                suggested_verb_replacements[weak_verb] = strong_verb

        # Keep missing keywords list concise
        suggested_missing_keywords = missing_keywords[:8]

        return {
            "ats_score": ats_score,
            "readability_score": readability_score,
            "formatting_score": formatting_score,
            "keyword_score": keyword_score,
            "project_score": project_score,
            "suggestions": {
                "missing_keywords": [kw.title() for kw in suggested_missing_keywords],
                "weak_bullet_points": weak_bullets[:5],
                "action_verb_replacements": suggested_verb_replacements,
                "quantifiable_suggestions": [
                    "Change 'Responsible for writing APIs' to 'Engineered 15+ REST APIs reducing response times by 30% using FastAPI.'",
                    "Change 'Helped team build the frontend' to 'Collaborated with a cross-functional team of 5 to design and implement a React frontend, increasing user conversion rate by 12%.'",
                    "Change 'Worked on CI/CD pipelines' to 'Orchestrated automated Jenkins/Docker pipelines, saving 4 hours of deployment overhead per week.'"
                ]
            }
        }

    @classmethod
    def generate_ats_resume(cls, text: str, template: str, format_type: str = "docx") -> bytes:
        """
        Generates a clean template-based Word document (.docx) or returns simple PDF bytes.
        """
        # Parse sections from existing resume text to construct the document
        lines = text.split("\n")
        lines = [line.strip() for line in lines if line.strip()]

        # Try to guess Name and Contact info
        name = "Professional Candidate"
        contact_info = "email@example.com | +1 (555) 019-2834 | github.com/candidate | linkedin.com/in/candidate"
        
        # Simple extraction heuristics
        for line in lines[:4]:
            if "@" in line:
                contact_info = line
            elif len(line) < 30 and not any(kw in line.lower() for kw in ["education", "skills", "experience", "resume", "curriculum"]):
                name = line

        # Split remaining sections
        sections = {"Summary": [], "Skills": [], "Experience": [], "Projects": [], "Education": []}
        current_section = "Summary"
        
        for line in lines:
            line_lower = line.lower()
            if "skills" in line_lower or "technologies" in line_lower:
                current_section = "Skills"
                continue
            elif "experience" in line_lower or "employment" in line_lower or "work history" in line_lower:
                current_section = "Experience"
                continue
            elif "projects" in line_lower:
                current_section = "Projects"
                continue
            elif "education" in line_lower or "academic" in line_lower:
                current_section = "Education"
                continue
            
            if line != name and line != contact_info:
                sections[current_section].append(line)

        # Style options
        primary_color = RGBColor(37, 99, 235) # Default Stripe Blue #2563EB
        text_color = RGBColor(15, 23, 42) # #0F172A
        muted_color = RGBColor(100, 116, 139) # #64748B
        
        if template == "minimal":
            primary_color = RGBColor(0, 0, 0)
            text_color = RGBColor(30, 41, 59)
            muted_color = RGBColor(100, 116, 139)
        elif template == "modern":
            primary_color = RGBColor(79, 70, 229) # Indigo #4F46E5
            text_color = RGBColor(15, 23, 42)
            muted_color = RGBColor(100, 116, 139)

        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Base style fonts
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        font.color.rgb = text_color

        # Header - Name
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.color.rgb = primary_color

        # Header - Contact Info
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_p.add_run(contact_info)
        contact_run.font.size = Pt(9.5)
        contact_run.font.color.rgb = muted_color

        doc.add_paragraph() # Spacer

        # Populate Sections
        for sec_name, content in sections.items():
            if not content:
                continue
            
            # Section Heading
            head_p = doc.add_paragraph()
            head_run = head_p.add_run(sec_name.upper())
            head_run.bold = True
            head_run.font.size = Pt(12)
            head_run.font.color.rgb = primary_color
            
            # Bottom border / divider line
            # Instead of XML manipulation, we use a simple text divider
            div_p = doc.add_paragraph()
            div_p.paragraph_format.space_after = Pt(4)
            div_run = div_p.add_run("―" * 58)
            div_run.font.size = Pt(6)
            div_run.font.color.rgb = muted_color

            # Section Content
            if sec_name == "Skills":
                skills_text = ", ".join(content)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.add_run(skills_text)
            else:
                for line in content:
                    p = doc.add_paragraph()
                    if line.startswith("-") or line.startswith("•") or line.startswith("*"):
                        p.paragraph_format.left_indent = Inches(0.4)
                        p.style = 'List Bullet'
                        p.add_run(line.lstrip("-•* "))
                    else:
                        p.paragraph_format.left_indent = Inches(0.2)
                        p.add_run(line)

            doc.add_paragraph() # Spacer

        # Return file bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        docx_bytes = file_stream.read()

        if format_type == "docx":
            return docx_bytes
        
        # If format_type == "pdf", let's create a minimal PDF representation using raw PDF syntax.
        # This is fully standard, standalone, and doesn't require third-party compiled software.
        # A simple PDF structure includes: Catalog, Pages, Page, Content Stream.
        pdf_stream = io.BytesIO()
        
        # Define pdf elements
        pdf_text = text.replace("(", "\\(").replace(")", "\\)")
        pdf_lines = pdf_text.split("\n")
        
        # Build raw text content stream
        content_lines = []
        content_lines.append("BT")
        content_lines.append("/F1 10 Tf")
        content_lines.append("12 TL")
        content_lines.append("72 720 Td")
        
        # We write up to 45 lines to fit single page representation
        line_count = 0
        for l in pdf_lines[:45]:
            # sanitize string
            l_clean = l.encode('ascii', errors='ignore').decode('ascii')
            content_lines.append(f"({l_clean}) Tj T*")
            line_count += 1
            
        content_lines.append("ET")
        content_data = "\n".join(content_lines).encode('ascii')
        
        # Create standard PDF layout objects
        # Object 1: Catalog
        # Object 2: Outlines
        # Object 3: Pages
        # Object 4: Page 1
        # Object 5: Content stream
        # Object 6: Font
        
        objects = {}
        # Object 1: Catalog
        objects[1] = b"<< /Type /Catalog /Pages 3 0 R >>"
        # Object 2: Outlines
        objects[2] = b"<< /Type /Outlines /Count 0 >>"
        # Object 3: Pages
        objects[3] = b"<< /Type /Pages /Kids [4 0 R] /Count 1 >>"
        # Object 4: Page 1
        objects[4] = b"<< /Type /Page /Parent 3 0 R /MediaBox [0 0 612 792] /Contents 5 0 R /Resources << /Font << /F1 6 0 R >> >> >>"
        # Object 5: Content stream
        objects[5] = f"<< /Length {len(content_data)} >>\nstream\n".encode('ascii') + content_data + b"\nendstream"
        # Object 6: Font
        objects[6] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
        
        # Write PDF header
        pdf_stream.write(b"%PDF-1.4\n")
        
        # Keep track of offsets
        offsets = {}
        for obj_id, obj_bytes in objects.items():
            offsets[obj_id] = pdf_stream.tell()
            pdf_stream.write(f"{obj_id} 0 obj\n".encode('ascii'))
            pdf_stream.write(obj_bytes)
            pdf_stream.write(b"\nendobj\n")
            
        # Write Cross-Reference Table
        xref_offset = pdf_stream.tell()
        pdf_stream.write(b"xref\n")
        pdf_stream.write(f"0 {len(objects) + 1}\n".encode('ascii'))
        pdf_stream.write(b"0000000000 65535 f \n")
        for obj_id in sorted(objects.keys()):
            offset_str = f"{offsets[obj_id]:010d} 00000 n \n"
            pdf_stream.write(offset_str.encode('ascii'))
            
        # Write trailer
        pdf_stream.write(b"trailer\n")
        pdf_stream.write(f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode('ascii'))
        pdf_stream.write(b"startxref\n")
        pdf_stream.write(f"{xref_offset}\n".encode('ascii'))
        pdf_stream.write(b"%%EOF\n")
        
        pdf_stream.seek(0)
        return pdf_stream.read()
